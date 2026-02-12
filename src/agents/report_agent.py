"""
Report Agent module for generating downloadable PDF and DOCX reports.
"""

from __future__ import annotations

import os
from datetime import datetime
import pandas as pd
from typing import Any

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config.settings import Settings
from src.utils.logger import setup_logger
from src.utils.error_handler import ReportGenerationError

logger = setup_logger("report_agent")


class ReportAgent:
    """Agent responsible for creating formatted reports."""

    def __init__(self):
        Settings.ensure_dirs()
        self.output_dir = Settings.REPORTS_DIR

    def _get_filename(self, prefix: str, ext: str) -> str:
        """Generate a timestamped filename."""
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_prefix = "".join([c if c.isalnum() else "_" for c in prefix[:20]])
        return os.path.join(self.output_dir, f"{safe_prefix}_{ts}.{ext}")

    def generate_pdf_report(self, data: dict) -> str:
        """Generate a PDF report from the analysis data."""
        try:
            filepath = self._get_filename("Report", "pdf")
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Custom styles
            title_style = styles['Title']
            heading_style = styles['Heading2']
            normal_style = styles['Normal']
            code_style = ParagraphStyle('Code', parent=styles['Normal'], fontName='Courier', fontSize=8)

            # 1. Header
            story.append(Paragraph("IntelliQuery AI Report", title_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"Query: {data.get('question', 'N/A')}", heading_style))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
            story.append(Spacer(1, 24))

            # 2. Insights
            if 'insights' in data:
                story.append(Paragraph("Executive Summary", heading_style))
                story.append(Paragraph(data['insights'].get('summary', ''), normal_style))
                story.append(Spacer(1, 12))
                
                story.append(Paragraph("Key Insights", heading_style))
                for insight in data['insights'].get('key_insights', []):
                    story.append(Paragraph(f"• {insight}", normal_style))
                story.append(Spacer(1, 12))

                story.append(Paragraph("Recommendations", heading_style))
                for rec in data['insights'].get('recommendations', []):
                    story.append(Paragraph(f"• {rec}", normal_style))
                story.append(Spacer(1, 24))

            # 3. Visualization
            # Note: In a real app we'd need to save the plotly fig to an image bytes IO first
            # Since plotly static export requires kaleidoscope (extra dep), we might skip or handle if bytes provided
            # For this simplified version we skip image embedding in PDF to avoid complex dep issues unless user provided bytes
            story.append(Paragraph("Visualization", heading_style))
            story.append(Paragraph("(Visualizations are available in the interactive dashboard)", normal_style))
            story.append(Spacer(1, 24))

            # 4. Data Preview
            if 'results' in data and isinstance(data['results'], pd.DataFrame):
                df = data['results'].head(20)
                story.append(Paragraph(f"Data Preview (Top {len(df)} rows)", heading_style))
                
                # Convert DF to list of lists for Table
                table_data = [df.columns.tolist()] + df.values.tolist()
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(t)
                story.append(Spacer(1, 24))

            # 5. SQL
            story.append(Paragraph("Generated SQL Query", heading_style))
            story.append(Paragraph(data.get('sql_query', ''), code_style))

            doc.build(story)
            logger.info(f"PDF Report generated: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise ReportGenerationError(f"Failed to generate PDF: {e}")

    def generate_docx_report(self, data: dict) -> str:
        """Generate a Word DOCX report from the analysis data."""
        try:
            filepath = self._get_filename("Report", "docx")
            doc = Document()

            # 1. Header
            doc.add_heading('IntelliQuery AI Report', 0)
            p = doc.add_paragraph()
            p.add_run(f"Query: {data.get('question', 'N/A')}").bold = True
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

            # 2. Insights
            if 'insights' in data:
                doc.add_heading('Executive Summary', level=1)
                doc.add_paragraph(data['insights'].get('summary', ''))
                
                doc.add_heading('Key Insights', level=2)
                for insight in data['insights'].get('key_insights', []):
                    doc.add_paragraph(insight, style='List Bullet')
                
                doc.add_heading('Recommendations', level=2)
                for rec in data['insights'].get('recommendations', []):
                    doc.add_paragraph(rec, style='List Bullet')

            # 3. Data Preview
            if 'results' in data and isinstance(data['results'], pd.DataFrame):
                df = data['results'].head(20)
                doc.add_heading(f'Data Preview (Top {len(df)} rows)', level=1)
                
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                
                # Rows
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, item in enumerate(row):
                         row_cells[i].text = str(item)

            # 4. SQL
            doc.add_heading('Generated SQL Query', level=1)
            p = doc.add_paragraph(data.get('sql_query', ''))
            p.style = 'Quote'

            doc.save(filepath)
            logger.info(f"DOCX Report generated: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise ReportGenerationError(f"Failed to generate DOCX: {e}")
